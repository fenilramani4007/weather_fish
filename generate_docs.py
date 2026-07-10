"""
Generates WEATHER-FISH project documentation as a Word (.docx) file.
Run: python generate_docs.py
Output: WEATHER-FISH_Documentation.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_height       = Cm(29.7)
section.page_width        = Cm(21.0)
section.left_margin       = Cm(2.5)
section.right_margin      = Cm(2.5)
section.top_margin        = Cm(2.5)
section.bottom_margin     = Cm(2.5)

# ── Style helpers ─────────────────────────────────────────────────────────────
DARK_BLUE = RGBColor(0x0D, 0x1B, 0x3E)
GOLD      = RGBColor(0xC9, 0xA2, 0x27)
GRAY      = RGBColor(0x44, 0x44, 0x44)
CODE_BG   = RGBColor(0xF4, 0xF4, 0xF4)


def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = DARK_BLUE
    p.runs[0].font.size = Pt(20)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = GOLD
    p.runs[0].font.size = Pt(14)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = GRAY
    p.runs[0].font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    return p


def body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(3)
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    return p


def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0x6B)
    # Light gray shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F0F0F0")
    pPr.append(shd)
    return p


def qa(q, a):
    p = doc.add_paragraph()
    r1 = p.add_run("Q: " + q)
    r1.bold = True
    r1.font.color.rgb = DARK_BLUE
    r1.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)

    p2 = doc.add_paragraph()
    r2 = p2.add_run("A: " + a)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    p2.paragraph_format.left_indent = Inches(0.2)
    p2.paragraph_format.space_after = Pt(6)


def divider():
    p = doc.add_paragraph("─" * 80)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.runs[0].font.size = Pt(8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("\n\n🐟  WEATHER-FISH")
r.font.size = Pt(32)
r.bold = True
r.font.color.rgb = DARK_BLUE

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = tp2.add_run("KI-gestützte adaptive Wetterberichte")
r2.font.size = Pt(16)
r2.font.color.rgb = GOLD
r2.italic = True

tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tp3.add_run(
    "\nComplete Project Documentation\n"
    "Step-by-Step Guide · Code Examples · Q&A\n\n"
    "OTH Amberg-Weiden — PMAE Course — 2025/26\n"
    "Student: Fenil Ramani\n"
)
r3.font.size = Pt(12)
r3.font.color.rgb = GRAY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════
h1("Table of Contents")
toc_items = [
    ("1", "Project Overview"),
    ("2", "System Architecture"),
    ("3", "Technology Stack"),
    ("4", "Step-by-Step Setup Guide"),
    ("   4.1", "Prerequisites"),
    ("   4.2", "Local Development Setup"),
    ("   4.3", "Environment Variables"),
    ("   4.4", "Running Locally"),
    ("   4.5", "Docker Build"),
    ("   4.6", "Hugging Face Spaces Deployment"),
    ("5", "Core Features — How They Work"),
    ("   5.1", "Weather Data Pipeline"),
    ("   5.2", "AI Text Generation (Gemini)"),
    ("   5.3", "Persona System"),
    ("   5.4", "Text-to-Speech (TTS)"),
    ("   5.5", "AI Chat Assistant"),
    ("   5.6", "Language Toggle (DE / EN)"),
    ("   5.7", "PLZ + City Autocomplete Search"),
    ("   5.8", "Saved Locations"),
    ("   5.9", "MongoDB Storage"),
    ("   5.10", "Caching"),
    ("   5.11", "Hourly Auto-Generation (Scheduler)"),
    ("6", "API Reference"),
    ("7", "Frontend Component Guide"),
    ("8", "Security & Best Practices"),
    ("9", "Common Errors & Solutions"),
    ("10", "Questions & Answers"),
    ("11", "AI Tool Declaration"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{num:<8}{title}")
    r.font.size = Pt(11)
    r.font.name = "Courier New"

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Project Overview")

body(
    "WEATHER-FISH is an AI-driven adaptive weather narration system developed as part of "
    "the PMAE (Praktische Methoden der Angewandten Elektronik) course at OTH Amberg-Weiden. "
    "The system fetches real-time weather data for German postal codes and generates spoken "
    "weather reports in three distinct AI personas, with text-to-speech audio, a conversational "
    "chat assistant, and support for both German and English languages."
)

h2("Core Goals (from Project Definition)")
bullet("Fetch and process real-time weather data using OpenWeatherMap API")
bullet("Generate personalized, broadcast-style narrations using Google Gemini AI")
bullet("Present reports through three distinct persona voices (Fisch, Merkel, Haftbefehl)")
bullet("Convert generated text to speech using Edge-TTS")
bullet("Store and retrieve data persistently using MongoDB Atlas")
bullet("Provide a modern React-based web interface with German/English language support")
bullet("Deploy as a containerized application on Hugging Face Spaces")
bullet("Include an AI chat assistant grounded in live weather data")

h2("What Makes It Adaptive")
body(
    "The system is 'adaptive' because the AI narrator adjusts its tone based on weather severity. "
    "The Context Engine (context_engine.py) analyzes live data and classifies the situation as: "
    "normal, warning (severe weather), or cheerful (sunny). This classification then shapes the "
    "Gemini prompt — instructing the model to be calm, urgent, or upbeat respectively."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
h1("2. System Architecture")

body(
    "WEATHER-FISH follows a layered architecture with clear separation between data "
    "acquisition, AI processing, storage, and presentation."
)

h2("Architecture Diagram (text representation)")
code_block(
    "┌─────────────────────────────────────────────────────────────────┐\n"
    "│                        USER BROWSER                             │\n"
    "│   React 18 + TypeScript + Vite  (SPA on port 5173 in dev)       │\n"
    "│   Components: WeatherSidebar, CurrentWeather, TextReport,        │\n"
    "│               WeatherVisuals, SavedLocations, ChatWidget          │\n"
    "└──────────────────────┬──────────────────────────────────────────┘\n"
    "                       │ HTTP (same origin in prod / CORS in dev)\n"
    "┌──────────────────────▼──────────────────────────────────────────┐\n"
    "│                FastAPI Backend  (port 7860 / 8000)               │\n"
    "│  main.py → api.py → data.py → context_engine.py                 │\n"
    "│                            → text_generation.py → tts.py         │\n"
    "│                            → mongo.py (read/write)               │\n"
    "└──────┬────────────┬──────────────────┬────────────────┬─────────┘\n"
    "       │            │                  │                │\n"
    " OpenWeather    Google Gemini     Edge-TTS         MongoDB Atlas\n"
    "   API 3.0      (AI text gen)   (speech audio)    (persistence)\n"
)

h2("Request Flow — Full Pipeline")
numbered("User enters a German postal code (PLZ) in the sidebar")
numbered("Frontend sends POST /generate-documents with {zipcodes, cities, person, language}")
numbered("api.py calls data.py → fetches current, hourly, and 7-day forecast from OpenWeatherMap")
numbered("Weather JSON is stored in MongoDB (weather_snapshots collection)")
numbered("context_engine.py analyzes severity, time of day, temperature feel")
numbered("text_generation.py builds a Gemini prompt including persona, weather data, and context")
numbered("Gemini generates 6–8 sentence narration; result stored in MongoDB (reports collection)")
numbered("tts.py converts text to MP3 audio using Edge-TTS (saved to /data/speech/)")
numbered("Frontend polls /api/weather/{zipcode}, /api/report/{presenter}, /speech/{presenter}.mp3")
numbered("UI displays weather cards, plays audio, shows text report")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Technology Stack")

h2("Backend")
bullet("Python 3.11 — runtime")
bullet("FastAPI — REST API framework with async support")
bullet("Uvicorn — ASGI server (port 7860 on HF Spaces, 8000 locally)")
bullet("google-genai 2.8.0 — Gemini AI SDK (text generation + chat)")
bullet("edge-tts — Microsoft Edge TTS voices for German and English")
bullet("motor / pymongo — MongoDB async/sync driver")
bullet("APScheduler — background hourly job scheduler")
bullet("requests-cache — persistent HTTP caching for OpenWeatherMap calls")
bullet("python-dotenv — loads .env file for local development")

h2("Frontend")
bullet("React 18 + TypeScript — UI framework")
bullet("Vite 6 — build tool and dev server")
bullet("React Context API — shared state (LanguageContext, LocationContext, WeatherContext)")
bullet("Inline CSS / custom properties — dark theme with gold accent (--gold: #c9a227)")

h2("Infrastructure")
bullet("Docker — containerized deployment (Dockerfile in repo root)")
bullet("Hugging Face Spaces — free cloud hosting (SDK: docker)")
bullet("MongoDB Atlas M0 — free tier cloud database")
bullet("Git LFS — stores PNG mascot images without bloating git history")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. STEP-BY-STEP SETUP GUIDE
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Step-by-Step Setup Guide")

# 4.1
h2("4.1 Prerequisites")
body("Install the following before starting:")
bullet("Python 3.11+ — python.org/downloads")
bullet("Node.js 20+ and npm — nodejs.org")
bullet("Git — git-scm.com")
bullet("Docker Desktop (for container testing) — docker.com")
bullet("A free MongoDB Atlas account — mongodb.com/atlas")
bullet("A free Google AI Studio API key — aistudio.google.com")
bullet("A free OpenWeatherMap API key — openweathermap.org/api")

# 4.2
h2("4.2 Local Development Setup")
numbered("Clone the repository:")
code_block("git clone <your-repo-url>\ncd weatherfish")

numbered("Create a Python virtual environment:")
code_block("cd src\npython -m venv .venv\n.venv\\Scripts\\activate        # Windows\n# source .venv/bin/activate   # Mac/Linux")

numbered("Install Python dependencies:")
code_block("pip install -r requirements.txt")

numbered("Install frontend dependencies:")
code_block("cd frontend\nnpm install\ncd ..")

# 4.3
h2("4.3 Environment Variables")
body(
    "Create the file src/.env with these values. This file must NEVER be committed to git — "
    "it is listed in .gitignore."
)
code_block(
    "GEMINI_API_KEY=your_gemini_key_here\n"
    "OPENWEATHER_API_KEY=your_openweather_key_here\n"
    "MONGODB_URI=mongodb+srv://username:password@cluster.mongodb.net/weatherfish?retryWrites=true&w=majority"
)
body(
    "IMPORTANT: If your MongoDB password contains a # character, you must URL-encode it as %23 "
    "in the URI string. Example: password FPS#76852 becomes FPS%2376852 in the URI."
)
code_block(
    "# Correct:\n"
    "MONGODB_URI=mongodb+srv://user:FPS%2376852@cluster.mongodb.net/weatherfish?...\n\n"
    "# WRONG — will break URI parsing:\n"
    "MONGODB_URI=mongodb+srv://user:FPS#76852@cluster.mongodb.net/weatherfish?..."
)

# 4.4
h2("4.4 Running Locally")
body("You need two terminals running simultaneously:")

body("Terminal 1 — Backend:")
code_block(
    "cd src\n"
    ".venv\\Scripts\\activate\n"
    "python -m uvicorn backend.main:app --reload --port 8000"
)

body("Terminal 2 — Frontend (Vite dev server):")
code_block(
    "cd src/frontend\n"
    "npm run dev"
)

body("Open http://localhost:5173 in your browser. The frontend dev server proxies API calls to port 8000.")

# 4.5
h2("4.5 Docker Build (local test)")
numbered("Build the frontend first:")
code_block("cd src/frontend && npm run build && cd ../..")

numbered("Build the Docker image:")
code_block("docker build -t weatherfish .")

numbered("Run the container:")
code_block(
    "docker run -p 7860:7860 \\\n"
    "  -e GEMINI_API_KEY=your_key \\\n"
    "  -e OPENWEATHER_API_KEY=your_key \\\n"
    "  -e MONGODB_URI=\"your_uri\" \\\n"
    "  weatherfish"
)

numbered("Open http://localhost:7860")

# 4.6
h2("4.6 Hugging Face Spaces Deployment")
numbered("Create a new Space at huggingface.co/new-space")
numbered("Select SDK: Docker and make the Space public")
numbered("Add secrets in Space Settings → Variables and Secrets:")
bullet("GEMINI_API_KEY", level=1)
bullet("OPENWEATHER_API_KEY", level=1)
bullet("MONGODB_URI (with URL-encoded password)", level=1)
numbered("Push your code to the Space remote:")
code_block(
    "git remote add hf https://huggingface.co/spaces/YOUR_USERNAME/weather-fish\n"
    "git push hf main"
)
numbered("HF builds the Docker image automatically. Monitor progress in the Logs tab.")
numbered("Access at: https://YOUR_USERNAME-weather-fish.hf.space")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. CORE FEATURES
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Core Features — How They Work")

# 5.1
h2("5.1 Weather Data Pipeline")
body(
    "The data pipeline is orchestrated by src/backend/api.py. It calls OpenWeatherMap's "
    "One Call API 3.0 to fetch three data sets for each location:"
)
bullet("Current conditions — temperature, feels-like, humidity, wind speed, sky condition")
bullet("Hourly forecast — next 48 hours broken into hour-by-hour entries")
bullet("Daily forecast (daily_weekone) — 7-day outlook with min/max temp, rain probability, wind")

body("Example of a weather document stored in MongoDB:")
code_block(
    '{\n'
    '  "zipcode": "90402",\n'
    '  "city": "Nürnberg",\n'
    '  "current": {\n'
    '    "temperature": 22,\n'
    '    "feels like": 21,\n'
    '    "humidity": 58,\n'
    '    "overcast": "partly cloudy",\n'
    '    "current_precipitation": null\n'
    '  },\n'
    '  "hourly": {\n'
    '    "9": { "temperature": 19, "precipitation probability": 10 },\n'
    '    "14": { "temperature": 24, "precipitation probability": 5 }\n'
    '  },\n'
    '  "daily_weekone": {\n'
    '    "2026-06-14": { "mintemp": 16, "maxtemp": 26, "overcast": "sunny" }\n'
    '  },\n'
    '  "fetched_at": "2026-06-13T10:00:00Z"\n'
    '}'
)

# 5.2
h2("5.2 AI Text Generation (Gemini)")
body(
    "File: src/ai/text_generation.py"
    "\n\nThe system uses Google Gemini to generate broadcast-style weather narrations. "
    "The process has two stages: building a rich prompt, then calling the Gemini API."
)

h3("Prompt Structure")
body("Each prompt is assembled from these parts in order:")
numbered("Persona instruction — defines who is speaking")
numbered("Weather data — structured JSON for all requested locations")
numbered("Context — severity (normal/warning/cheerful), time of day, temperature feel")
numbered("Highlights — key events to weave in (e.g., 'heavy rain likely at 14:00')")
numbered("Forecast instructions — use hourly for today's arc, daily for tomorrow")
numbered("Personalisation — if user provided hobbies, include an activity tip")
numbered("Language — ISO 639-1 code ('de' for German, 'en' for English)")
numbered("Format rules — prose only, 6–8 sentences, specific structure")

h3("Model Fallback Chain")
body(
    "To handle rate limits and quota exhaustion, the system tries models in this order, "
    "falling back automatically when a model's daily quota is exceeded:"
)
code_block(
    "MODELS = [\n"
    '    "gemini-1.5-flash",       # 1,500 req/day free\n'
    '    "gemini-1.5-flash-8b",    # 1,500 req/day, faster\n'
    '    "gemini-2.0-flash-lite",  # separate quota bucket\n'
    '    "gemini-2.0-flash",       # 200 req/day\n'
    "]"
)
body(
    "If ALL models fail (total quota exhausted), the system returns a template fallback — "
    "a pre-written German text with the live weather data inserted."
)

# 5.3
h2("5.3 Persona System")
body("Three presenter personas are built into the system:")

h3("Fisch (Fish Mascot)")
body(
    "The default persona. Cheerful, warm-hearted fish mascot who loves weather. "
    "Friendly and witty, occasionally uses subtle fish-themed humour. "
    "Voice: de-DE-KillianNeural (German male, warm tone)."
)
code_block('Example: "Hallo aus dem Aquarium! Das Wasser — äh, die Luft — riecht heute nach Sommer..."')

h3("Merkel")
body(
    "Angela Merkel style. Composed, measured, authoritative, data-driven. "
    "Occasional dry German wit. Every word counts. "
    "Voice: de-DE-AmalaNeural (German female, authoritative)."
)
code_block('Example: "Wir haben die Situation analysiert. Die Temperaturen liegen bei 22 Grad..."')

h3("Haftbefehl")
body(
    "German rapper from Offenbach. Punchy, energetic, urban slang, swagger and rhythm. "
    "Accurate information delivered with attitude. "
    "Voice: de-DE-BerndNeural (German male, energetic)."
)
code_block('Example: "Yo, Wetter-Check! 22 Grad, kein Regen, Sonne pusht — heute ist dein Tag, Bruder!"')

# 5.4
h2("5.4 Text-to-Speech (TTS)")
body(
    "File: src/ai/tts.py"
    "\n\nThe system converts generated text to MP3 audio using Microsoft Edge-TTS (edge-tts Python library). "
    "This is free and requires no API key."
)

h3("The Event Loop Problem and Solution")
body(
    "FastAPI runs on Uvicorn, which manages its own asyncio event loop. "
    "Edge-TTS is also async. Calling asyncio.run() from within an already-running "
    "event loop raises RuntimeError. The solution is to run TTS in a completely "
    "separate thread (ThreadPoolExecutor), where that thread can safely call asyncio.run():"
)
code_block(
    "def generate_mp3(language_code, text, person):\n"
    "    def _run_in_thread():\n"
    "        asyncio.run(_generate_async(text, voice, output_path))\n"
    "    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:\n"
    "        executor.submit(_run_in_thread).result(timeout=60)"
)

body("Audio files are saved to SPEECH_DIR:")
bullet("Local development: src/frontend/public/speech/")
bullet("HF Spaces: /data/speech/ (persistent volume — survives container restarts)")

# 5.5
h2("5.5 AI Chat Assistant")
body(
    "File: src/ai/chat.py — Backend\n"
    "File: src/frontend/src/components/ChatWidget.tsx — Frontend"
    "\n\nThe chat assistant is a floating 🐟 button (FAB) in the bottom-right corner. "
    "It uses Gemini 2.0 Flash with the user's current live weather data as grounding context."
)

h3("How Context Is Injected")
body(
    "When the user opens the chat widget, the frontend sends the current postal code. "
    "The backend retrieves the latest weather document from MongoDB and builds a system prompt "
    "that includes:"
)
bullet("Current temperature, feels-like, humidity, sky condition, precipitation")
bullet("Today's temperature range (min/max from hourly data)")
bullet("Hours when rain probability exceeds 60%")
bullet("Tomorrow's forecast (min/max temp, sky condition, wind speed)")

body("This gives Gemini accurate, location-specific context to answer questions like:")
bullet('"Do I need an umbrella today?"')
bullet('"What should I wear tomorrow morning?"')
bullet('"Is it good weather for cycling this afternoon?"')

h3("Conversation History")
body(
    "The last 10 messages (5 exchanges) are sent with each request, enabling "
    "multi-turn conversations. The system prompt is injected into the first user "
    "message so Gemini has context from the very start."
)

# 5.6
h2("5.6 Language Toggle (DE / EN)")
body(
    "Files: src/frontend/src/contexts/LanguageContext.tsx, Header.tsx, WeatherSidebar.tsx, ChatWidget.tsx"
    "\n\nThe language preference is managed through React Context API, making it globally "
    "reactive without prop drilling."
)

h3("How It Works")
numbered("LanguageContext stores the current language ('de' | 'en') in both React state and localStorage")
numbered("LanguageProvider wraps the entire app in App.tsx")
numbered("Any component calls useLanguage() to read or set the language")
numbered("When language changes, ALL components re-render with the new language instantly")
numbered("The preference persists across page reloads via localStorage")

h3("Backend Language Support")
body("The language parameter is sent in API requests and affects:")
bullet("Gemini text generation — prompt instructs which language to use")
bullet("TTS voice selection — German voices for 'de', English voices for 'en'")
bullet("Chat assistant — responds in the selected language")

# 5.7
h2("5.7 PLZ + City Autocomplete Search")
body(
    "File: src/frontend/src/components/WeatherSidebar.tsx"
    "\n\nInstead of requiring exact input, the search box offers an autocomplete dropdown "
    "that searches a local JSON file (public/postal_codes/postal_codes.json) with ~16,000 "
    "German postal codes."
)

h3("Search Logic")
bullet("Type digits (e.g., '903') → filters by postal code prefix")
bullet("Type letters (e.g., 'Nür') → filters by city name (case-insensitive)")
bullet("Results are limited to the top 8 matches for performance")
bullet("Selecting an entry auto-fills both PLZ and city name — no manual typing needed")
bullet("Dropdown closes automatically when user clicks outside (ref-based handler)")

h3("Why This Approach")
body(
    "Earlier versions required users to type the exact German city name, which failed for "
    "English spellings (e.g., 'Nuremberg' instead of 'Nürnberg') and umlauts. "
    "The autocomplete approach bypasses this entirely — users always select from valid entries."
)

# 5.8
h2("5.8 Saved Locations")
body(
    "File: src/frontend/src/components/SavedLocations.tsx"
    "\n\nUsers can save multiple locations. Each saved location is stored in MongoDB "
    "(locations collection) and appears as a clickable card below the search box. "
    "Clicking a saved location immediately loads its weather without re-entering the PLZ."
)
bullet("GET /api/locations — retrieve all saved locations")
bullet("DELETE /api/locations/{zipcode} — remove a saved location")
bullet("Locations are auto-saved when a user generates a weather report")

# 5.9
h2("5.9 MongoDB Storage")
body(
    "File: src/database/mongo.py"
    "\n\nMongoDB Atlas M0 (free tier) stores three collections:"
)
bullet("weather_snapshots — one document per zipcode, updated on each pipeline run")
bullet("reports — one document per presenter (Fisch/Merkel/Haftbefehl), with the AI-generated text")
bullet("locations — all user-saved locations with added_at and last_seen timestamps")

h3("Upsert Pattern")
body("All writes use upsert (update or insert) to avoid duplicates:")
code_block(
    "collection.update_one(\n"
    '    {"zipcode": zipcode},          # filter\n'
    "    {\"$set\": document},           # update\n"
    "    upsert=True                    # create if not exists\n"
    ")"
)

h3("Password Masking")
body(
    "The MongoDB URI password is never exposed in logs or API responses. "
    "The _masked_uri() function replaces the password with *** in all output:"
)
code_block(
    "# Input:  mongodb+srv://user:FPS%2376852@cluster.../weatherfish\n"
    "# Output: mongodb+srv://user:***@cluster.../weatherfish"
)

# 5.10
h2("5.10 Caching")
body(
    "File: src/database/data.py"
    "\n\nOpenWeatherMap API calls are cached using requests-cache. "
    "This prevents making redundant API calls when the same location "
    "is requested multiple times within the cache window."
)
bullet("Cache TTL: 1 hour (3600 seconds)")
bullet("Local development: cached in src/database/.cache (SQLite)")
bullet("HF Spaces: cached in /data/weather_cache (persistent volume, survives restarts)")

body("Cache path selection logic:")
code_block(
    '_speech_dir = os.environ.get("SPEECH_DIR", "")\n'
    "CACHE_PATH = (\n"
    "    os.path.join(os.path.dirname(_speech_dir), 'weather_cache')\n"
    "    if _speech_dir       # production: /data/weather_cache\n"
    "    else os.path.join(BASE_DIR, '.cache')  # development\n"
    ")"
)

# 5.11
h2("5.11 Hourly Auto-Generation (Scheduler)")
body(
    "File: src/backend/main.py — APScheduler"
    "\n\nThe system automatically regenerates weather reports every hour for all "
    "saved locations, so reports are always fresh without user interaction."
)
bullet("Scheduler starts at app startup (lifespan event)")
bullet("Job ID: 'hourly_generation', interval: 1 hour")
bullet("Fetches all saved locations from MongoDB")
bullet("Runs the full pipeline (data fetch → AI text → TTS) for each location")
bullet("Status available at GET /api/schedule/status")
bullet("Manual trigger available at POST /api/schedule/run")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. API REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
h1("6. API Reference")

body("Base URL (production): https://USERNAME-weather-fish.hf.space")
body("Base URL (development): http://localhost:8000")
body("Interactive docs: /docs (Swagger UI auto-generated by FastAPI)")

endpoints = [
    ("POST", "/generate-documents",
     "Run the full pipeline for one or more postal codes.",
     '{"cities": ["Nürnberg"], "zipcodes": ["90402"], "person": "Fisch", "language": "de"}',
     '{"status": "success", "message": "Wetterdaten wurden erzeugt."}'),

    ("GET", "/api/weather/{zipcode}",
     "Retrieve the latest weather snapshot for a postal code from MongoDB.",
     None,
     '{"zipcode": "90402", "city": "Nürnberg", "current": {...}, "hourly": {...}}'),

    ("GET", "/api/report/{presenter}",
     "Retrieve the latest AI-generated text report for a presenter.",
     None,
     '{"presenter": "Fisch", "text": "Hallo aus dem Aquarium! ...", "generated_at": "..."}'),

    ("GET", "/api/locations",
     "List all saved locations.",
     None,
     '{"locations": [...], "count": 3}'),

    ("DELETE", "/api/locations/{zipcode}",
     "Remove a saved location.",
     None,
     '{"status": "deleted", "zipcode": "90402"}'),

    ("GET", "/api/status",
     "Health check — MongoDB connectivity and data counts.",
     None,
     '{"status": "ok", "mongodb": "connected", "snapshots": 5, "reports": 3}'),

    ("POST", "/api/chat",
     "Send a message to the AI weather assistant.",
     '{"message": "Do I need an umbrella?", "zipcode": "90402", "history": [], "language": "de"}',
     '{"reply": "Nein, heute kein Regen. Sonnig mit 22°C."}'),

    ("GET", "/speech/{filename}",
     "Serve a TTS audio file (e.g., Fisch.mp3).",
     None,
     "audio/mpeg binary stream"),

    ("GET", "/api/schedule/status",
     "Check scheduler status and last run time.",
     None,
     '{"scheduler": "running", "next_run": "...", "last_status": "ok — 3 location(s)"}'),
]

for method, path, desc, req, resp in endpoints:
    h3(f"{method} {path}")
    body(desc)
    if req:
        body("Request body:")
        code_block(req)
    body("Response:")
    code_block(resp)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. FRONTEND COMPONENT GUIDE
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Frontend Component Guide")

components = [
    ("App.tsx", "Root component. Wraps everything in LanguageProvider, LocationProvider, WeatherProvider."),
    ("Layout.tsx", "Structural shell. Renders Header, sidebar (WeatherSidebar + SavedLocations), main content (CurrentWeather + TextReport + WeatherVisuals), and ChatWidget."),
    ("Header.tsx", "Top bar with WEATHER-FISH logo and a reactive language badge (KI-WETTER · DE or EN) that updates via useLanguage() context."),
    ("WeatherSidebar.tsx", "Main control panel. Contains the PLZ/city autocomplete search, persona selector, hobby input, language toggle, and the Generate button. Sends POST /generate-documents."),
    ("CurrentWeather.tsx", "Displays current conditions (temperature, feels-like, humidity, wind, sky icon) for the selected location."),
    ("TextReport.tsx", "Shows the AI-generated text report for the selected presenter. Includes audio playback controls."),
    ("WeatherVisuals.tsx", "Charts and visual cards: hourly temperature curve, precipitation probability, 7-day forecast overview."),
    ("SavedLocations.tsx", "List of pinned locations. Each card shows city name and last temperature. Click to switch location."),
    ("ChatWidget.tsx", "Floating 🐟 FAB chat button. Opens a 320×440px panel with message bubbles, input field, and Send button. Communicates with POST /api/chat."),
]

for name, desc in components:
    h3(name)
    body(desc)

h2("Context Providers")
bullet("LanguageContext — provides { language, setLanguage } to any component via useLanguage() hook")
bullet("LocationContext — provides { currentLocation, setCurrentLocation } — the active postal code/city")
bullet("WeatherContext — provides cached weather data to avoid redundant API calls")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. SECURITY & BEST PRACTICES
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Security & Best Practices")

h2("API Key Security")
bullet("All API keys are stored in src/.env — this file is in .gitignore and never committed")
bullet("On HF Spaces, keys are stored as Secrets (encrypted environment variables, not in code)")
bullet("The backend reads keys via os.environ['KEY_NAME'] — never hardcoded")
bullet("MongoDB passwords with special characters are URL-encoded (# → %23)")

h2("Password Masking in Logs")
body(
    "The MongoDB URI contains the database password. The _masked_uri() function ensures "
    "the password is never written to logs or returned in API responses:"
)
code_block(
    "# Log output:\n"
    "[MongoDB] Connected to mongodb+srv://weather_fish_oth:***@cluster.../weatherfish"
)

h2("Input Validation")
bullet("Audio file paths are validated: must end in .mp3, cannot contain / or .. (path traversal prevention)")
bullet("Chat messages require a non-empty 'message' field (HTTP 400 if missing)")
bullet("All MongoDB operations use parameterized queries (no string interpolation)")

h2("CORS Configuration")
body(
    "In development, CORS allows all origins (*). For production hardening, "
    "this should be restricted to your specific domain."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. COMMON ERRORS & SOLUTIONS
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Common Errors & Solutions")

errors = [
    (
        "MongoDB connection fails — AuthenticationFailed",
        "The password contains a special character (e.g., #) that breaks URI parsing.",
        "URL-encode the special character: # becomes %23. Example: FPS#76852 → FPS%2376852 in the URI."
    ),
    (
        "HF Space shows blank white page",
        "The Space is set to Private. JavaScript assets fail to load without authentication.",
        "Go to Space Settings and set visibility to Public."
    ),
    (
        "HF Space shows 'configuration error'",
        "README.md is missing the YAML frontmatter block that HF requires.",
        "Add this at the very top of README.md:\n---\ntitle: Weather Fish\nsdk: docker\npinned: false\n---"
    ),
    (
        "TTS audio not generated — RuntimeError event loop",
        "Calling asyncio.run() from within Uvicorn's running event loop raises RuntimeError.",
        "Run TTS in a ThreadPoolExecutor thread (separate thread has its own event loop)."
    ),
    (
        "Gemini returns error 429 (quota exceeded)",
        "The free tier has daily and per-minute limits.",
        "The system auto-falls back through 4 model tiers. If all fail, a template report is shown. Wait until the next day for quota reset."
    ),
    (
        "Git push rejected — Updates were rejected (fetch first)",
        "HF auto-created an initial commit (README) that your local branch doesn't have.",
        "Use: git push hf main --force  (only safe on first setup when HF repo is otherwise empty)"
    ),
    (
        "Binary files rejected by git push to HF",
        "HF enforces a 10MB file size limit. PNG mascot files exceed this.",
        "Use git-lfs: git lfs track '*.png' → git add .gitattributes → commit and push."
    ),
    (
        "src refspec main does not match any",
        "Your local branch is named 'master', not 'main'.",
        "Rename: git branch -m master main"
    ),
    (
        "Ungültige PLZ/Ort Kombination (Invalid postal code/city)",
        "Old validation matched user text against German city names — fails for English names.",
        "Fixed: autocomplete dropdown from postal_codes.json. User always selects valid entries."
    ),
    (
        "Audio plays but no sound / plays 0 seconds",
        "Speech directory is not persistent — files are lost on container restart.",
        "Set SPEECH_DIR=/data/speech as an HF Secret to use HF's persistent volume."
    ),
]

for error, cause, solution in errors:
    h3(f"Error: {error}")
    body(f"Cause: {cause}")
    body(f"Solution: {solution}")
    divider()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. QUESTIONS & ANSWERS
# ══════════════════════════════════════════════════════════════════════════════
h1("10. Questions & Answers")
body("Anticipated questions from evaluators, professors, or reviewers.")

h2("Architecture & Design")

qa(
    "Why did you choose FastAPI over Flask or Django?",
    "FastAPI provides native async support (required for concurrent weather API calls), "
    "automatic OpenAPI/Swagger documentation, and Pydantic-based type validation. "
    "Flask lacks async; Django is too heavyweight for a pure API backend. "
    "FastAPI's performance characteristics match well with Uvicorn's ASGI model."
)

qa(
    "Why MongoDB instead of a relational database like PostgreSQL?",
    "Weather data is inherently semi-structured — the number of hourly entries, forecast days, "
    "and nested fields varies per API response. MongoDB's document model stores this naturally "
    "without requiring a rigid schema or JOIN tables. The upsert pattern (update or insert) "
    "also makes replacing stale weather data very clean."
)

qa(
    "Why Hugging Face Spaces instead of Heroku or AWS?",
    "HF Spaces is free for public Docker deployments, has built-in git integration, "
    "persistent volume storage (/data), and is widely used in the AI community. "
    "It was the most practical free option for a student project with AI model usage."
)

qa(
    "How does the system stay up to date without users triggering generation?",
    "APScheduler runs a background job every hour that calls the full pipeline for all "
    "saved locations. This means weather data and AI reports are refreshed automatically "
    "as long as the app is running."
)

h2("AI & Machine Learning")

qa(
    "How does the system use AI?",
    "Three ways: (1) Google Gemini generates broadcast-style weather narration text from structured "
    "weather data. (2) Gemini powers the conversational chat assistant, grounded in live weather. "
    "(3) Microsoft Edge-TTS (neural voices) converts the generated text to speech audio."
)

qa(
    "What is 'adaptive' about the narration?",
    "The Context Engine (context_engine.py) analyzes the weather data before generating text. "
    "It classifies the situation as 'warning' (storm, heavy rain), 'cheerful' (sunny, warm), "
    "or 'normal', and derives the time of day. These classifications are included in the Gemini "
    "prompt as tone instructions — so the same data produces a calmer report on a cloudy day "
    "versus an energetic one on a sunny day."
)

qa(
    "What happens when the Gemini quota is exhausted?",
    "The system attempts 4 different Gemini models in priority order (gemini-1.5-flash first). "
    "Each model gets 3 retry attempts for per-minute rate limits. If all models fail due to "
    "daily quota exhaustion, a deterministic template fallback generates a report from the "
    "weather data without AI. Users see a note that AI reports are temporarily unavailable."
)

qa(
    "How does the chat assistant know about the weather?",
    "When the user sends a message, the backend fetches the most recent weather document "
    "for their selected location from MongoDB. This data is formatted into a system prompt "
    "injected into the conversation context before sending to Gemini. Gemini then answers "
    "questions about weather it has never directly accessed — the app provides the facts."
)

qa(
    "Why use Edge-TTS instead of Google Cloud TTS or AWS Polly?",
    "Edge-TTS is free and requires no API key — it uses the same voice synthesis as "
    "Microsoft Edge browser. The voice quality is very good (neural voices), it supports "
    "German natively, and it removes a billing dependency for the student project."
)

h2("Frontend & UX")

qa(
    "Why React instead of plain HTML/JavaScript?",
    "The application has complex shared state: language selection affects Header, WeatherSidebar, "
    "and ChatWidget simultaneously. Location selection affects CurrentWeather, TextReport, and "
    "WeatherVisuals. React's Context API manages this cleanly. Plain JS would require manual "
    "DOM synchronization across components, which is error-prone."
)

qa(
    "How does the language toggle work across components?",
    "LanguageContext stores a 'de' | 'en' value in both React state and localStorage. "
    "Any component that calls useLanguage() re-renders automatically when the language changes. "
    "localStorage ensures the preference survives page reloads."
)

qa(
    "Why was city name validation removed in favor of autocomplete?",
    "The original approach matched user text against German city names in postal_codes.json. "
    "This failed for English spellings (Nuremberg ≠ Nürnberg) and umlauts. "
    "The autocomplete dropdown always shows valid options from the JSON — the user selects "
    "rather than types the exact name."
)

h2("Deployment & Operations")

qa(
    "Why must the MongoDB password be URL-encoded?",
    "The MongoDB connection URI follows the RFC 3986 URI standard. In a URI, # is a "
    "reserved delimiter (marks the fragment). If a password contains #, the URI parser "
    "treats everything after it as a URL fragment and ignores it, causing authentication "
    "to fail. Encoding # as %23 tells the parser it is a literal character in the password."
)

qa(
    "How are secrets kept out of the git repository?",
    "The src/.env file is listed in .gitignore and has never been committed. "
    "On HF Spaces, secrets are added via Settings → Variables and Secrets (AES-encrypted). "
    "The code reads keys via os.environ['KEY_NAME'] — there are no hardcoded credentials anywhere."
)

qa(
    "Why does the HF Space need to be public?",
    "HF private Spaces authenticate the initial HTML page load via a signed URL token. "
    "However, subsequent requests for /assets/*.js files fail authentication, causing React "
    "to never mount. Making the Space public allows all assets to load normally."
)

qa(
    "What is the SPEECH_DIR environment variable for?",
    "On HF Spaces, the container filesystem is ephemeral — files written inside the container "
    "are lost when the container restarts. HF provides a persistent volume mounted at /data. "
    "Setting SPEECH_DIR=/data/speech ensures generated audio files survive restarts."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. AI TOOL DECLARATION
# ══════════════════════════════════════════════════════════════════════════════
h1("11. Declaration of AI Tool Usage")

body(
    "In accordance with the OTH Amberg-Weiden AI usage policy (Regelungen zur Nutzung von "
    "KI-Tools, valid from 2024 onwards), the following AI tools were used in the development "
    "of this project:"
)

h2("Tools Used")
bullet("Claude Code (Anthropic) — AI programming assistant")
bullet("Google Gemini (via google-genai SDK) — integrated into the product itself for text generation and chat")
bullet("Microsoft Edge-TTS — integrated into the product for text-to-speech synthesis")

h2("Scope of AI Assistance (Claude Code)")
body(
    "Claude Code was used as a development assistant throughout the implementation phase. "
    "Specifically, it assisted with:"
)
bullet("Code generation: FastAPI endpoint implementations, React component scaffolding")
bullet("Debugging: identifying the asyncio event loop conflict in TTS, URL-encoding issues in MongoDB URI")
bullet("Deployment: Dockerfile configuration, Hugging Face Spaces setup, git-lfs configuration")
bullet("Code review: security checks (path traversal, credential exposure)")

h2("Student Authorship")
body(
    "The following were authored and owned by the student (Fenil Ramani) independently:"
)
bullet("System architecture decisions and technology selection")
bullet("Project requirements analysis against the PMAE course specification")
bullet("Integration choices (which APIs to use, data model design)")
bullet("Testing and validation of all features on the live deployment")
bullet("This documentation and all written academic submissions")

h2("Compliance Statement")
body(
    "This use of AI tools falls within the permitted category of 'Code snippet generation "
    "or debugging help (if relevant to the subject)' as defined in the OTH AI policy. "
    "All AI-generated content has been reviewed, understood, and validated by the student. "
    "The core analysis, architecture decisions, and academic assessment remain the student's "
    "own work."
)

divider()

# Footer
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run(
    "\nWEATHER-FISH Documentation  |  Fenil Ramani  |  OTH Amberg-Weiden  |  2025/26\n"
    "Generated: June 2026"
)
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = r"d:\weatherfish\WEATHER-FISH_Documentation.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
